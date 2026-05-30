"""
Generate Group 4 (S)TeemBang testing deliverables as .docx / .xlsx.
Run: python generate_documents.py
Output: same folder as this script (documents/)
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parent
DATA = ROOT / "data"
PREFIX = "Group4_STeemBang"
PROJECT = "(S)TeemBang Fitness Tracking Application"
GROUP = "Group 4"
VERSION = "1.0"
DOC_DATE = "May 13, 2026"
MEMBERS = [
    "APAS, Christel Grace",
    "LABE, Sherwin",
    "TUAZON, John Rae",
    "ZERRUDO, Jase Karl",
]


def _load(name: str) -> dict:
    return json.loads((DATA / name).read_text(encoding="utf-8"))


def _title(doc: Document, text: str, level: int = 0) -> None:
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def _meta_block(doc: Document, identifier: str) -> None:
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Document ID", identifier),
        ("Version", VERSION),
        ("Date", DOC_DATE),
        ("Project", PROJECT),
        ("Group", GROUP),
        ("Authors", ", ".join(MEMBERS)),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = str(val)
    doc.add_paragraph()


def build_test_plan() -> None:
    d = _load("01_test_plan.json")
    doc = Document()
    _title(doc, d["title"])
    _meta_block(doc, d["identifier"])
    for section in d["sections"]:
        _title(doc, section["heading"], level=1)
        if "paragraphs" in section:
            for p in section["paragraphs"]:
                _para(doc, p)
        if "bullets" in section:
            _bullets(doc, section["bullets"])
        if "table" in section:
            _table(doc, section["table"]["headers"], section["table"]["rows"])
    _title(doc, "Approvals", level=1)
    _table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [[r, "", "", DOC_DATE] for r in ["Project Lead", "QA Lead", "Instructor / Evaluator"]],
    )
    doc.save(ROOT / f"{PREFIX}_Doc1_TestPlan.docx")


def _style_xlsx_header(ws, row: int, cols: int) -> None:
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(color="FFFFFF", bold=True)
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(wrap_text=True, vertical="top")


def build_test_cases() -> None:
    d = _load("02_test_cases_document.json")
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"
    headers = d["headers"]
    ws.append(headers)
    _style_xlsx_header(ws, 1, len(headers))
    for row in d["rows"]:
        ws.append(row)
    widths = [12, 14, 28, 22, 36, 18, 28, 28, 10, 10, 16, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws2 = wb.create_sheet("Execution Log")
    ws2.append(d["execution_log"]["headers"])
    _style_xlsx_header(ws2, 1, len(d["execution_log"]["headers"]))
    for row in d["execution_log"]["rows"]:
        ws2.append(row)
    for i, w in enumerate([12, 10, 14, 10, 14, 12], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    wb.save(ROOT / f"{PREFIX}_Doc2_TestCases.xlsx")


def build_system_test_report() -> None:
    d = _load("03_system_test_report.json")
    doc = Document()
    _title(doc, d["title"])
    _meta_block(doc, d["identifier"])
    for section in d["sections"]:
        _title(doc, section["heading"], level=1)
        if "paragraphs" in section:
            for p in section["paragraphs"]:
                _para(doc, p)
        if "bullets" in section:
            _bullets(doc, section["bullets"])
        if "table" in section:
            _table(doc, section["table"]["headers"], section["table"]["rows"])
    doc.save(ROOT / f"{PREFIX}_Doc5_SystemTestReport.docx")


def build_uat_plan_report() -> None:
    d = _load("04_uat_plan_and_report.json")
    doc = Document()
    _title(doc, d["title"])
    _meta_block(doc, d["identifier"])
    for section in d["sections"]:
        _title(doc, section["heading"], level=1)
        if "paragraphs" in section:
            for p in section["paragraphs"]:
                _para(doc, p)
        if "bullets" in section:
            _bullets(doc, section["bullets"])
        if "numbered" in section:
            _numbered(doc, section["numbered"])
        if "table" in section:
            _table(doc, section["table"]["headers"], section["table"]["rows"])
    if "signoff" in d:
        _title(doc, "Sign-Off", level=1)
        _table(doc, d["signoff"]["headers"], d["signoff"]["rows"])
    doc.save(ROOT / f"{PREFIX}_Doc6_UATPlanReport.docx")


def build_defect_log() -> None:
    d = _load("05_defect_bug_report_log.json")
    wb = Workbook()
    ws = wb.active
    ws.title = "Defect Log"
    headers = d["headers"]
    ws.append(headers)
    _style_xlsx_header(ws, 1, len(headers))
    for row in d["rows"]:
        ws.append(row)
    widths = [10, 24, 14, 10, 10, 14, 30, 24, 24, 12, 16, 12, 12, 20]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = min(w, 40)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(ROOT / f"{PREFIX}_Doc7_DefectLog.xlsx")


def build_regression_report() -> None:
    d = _load("06_regression_test_report.json")
    doc = Document()
    _title(doc, d["title"])
    _meta_block(doc, d["identifier"])
    for section in d["sections"]:
        _title(doc, section["heading"], level=1)
        if "paragraphs" in section:
            for p in section["paragraphs"]:
                _para(doc, p)
        if "table" in section:
            _table(doc, section["table"]["headers"], section["table"]["rows"])
    doc.save(ROOT / f"{PREFIX}_Doc8_RegressionTestReport.docx")


def build_summary_report() -> None:
    d = _load("07_test_summary_report.json")
    doc = Document()
    _title(doc, d["title"])
    _meta_block(doc, d["identifier"])
    for section in d["sections"]:
        _title(doc, section["heading"], level=1)
        if "paragraphs" in section:
            for p in section["paragraphs"]:
                _para(doc, p)
        if "bullets" in section:
            _bullets(doc, section["bullets"])
        if "table" in section:
            _table(doc, section["table"]["headers"], section["table"]["rows"])
    if "signoff" in d:
        _title(doc, "Sign-Off", level=1)
        _table(doc, d["signoff"]["headers"], d["signoff"]["rows"])
    doc.save(ROOT / f"{PREFIX}_Doc9_TestSummaryReport.docx")


def main() -> None:
    DATA.mkdir(parents=True, exist_ok=True)
    builders = [
        build_test_plan,
        build_test_cases,
        build_system_test_report,
        build_uat_plan_report,
        build_defect_log,
        build_regression_report,
        build_summary_report,
    ]
    for fn in builders:
        name = fn.__name__
        print(f"Generating {name}...")
        fn()
    print(f"Done — {len(builders)} files written to {ROOT}")


if __name__ == "__main__":
    main()
